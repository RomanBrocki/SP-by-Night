from __future__ import annotations

import json
import os
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = Path(os.environ.get("SPBN_DOCX", str(ROOT / "Sao_Paulo_by_Night.docx")))
SRC = ROOT / "tools" / "sp_by_night_source.json"
REPORT = ROOT / "tools" / "docx_index_coteries_import_report.md"


def norm(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace('"', "").replace("'", "")
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    s = re.sub(r"\s+", " ", s)
    return s


def dequote_name(s: str) -> str:
    s = s.replace("“", '"').replace("”", '"').replace("’", "'").replace("`", "'")
    return s


@dataclass
class IndexRow:
    tier: str
    name: str
    sect: str
    clan: str
    role: str
    domain: str
    vínculo: str
    from_section: str


def norm_sect(v: str) -> str:
    n = norm(v)
    if "camar" in n:
        return "Camarilla"
    if "anar" in n:
        return "Anarch"
    if "autar" in n:
        return "Autarquicos"
    if "indep" in n:
        return "Independentes"
    if "segunda inqu" in n:
        return "Segunda Inquisicao"
    if n == "mortal":
        return "Mortal"
    return v.strip()


def norm_clan(v: str) -> str:
    v = v.strip()
    aliases = {
        "Banu Haqim": "Banu Haqim",
        "Brujah": "Brujah",
        "Caitiff": "Caitiff",
        "Gangrel": "Gangrel",
        "Hecata": "Hecata",
        "Lasombra": "Lasombra",
        "Malkavian": "Malkavian",
        "Ministry": "Ministry",
        "Nosferatu": "Nosferatu",
        "Ravnos": "Ravnos",
        "Salubri": "Salubri",
        "Toreador": "Toreador",
        "Tremere": "Tremere",
        "Tzimisce": "Tzimisce",
        "Ventrue": "Ventrue",
        "Thin-Blood": "Thin Blood",
        "Thin Blood": "Thin Blood",
        "Mortal": "Mortal",
    }
    return aliases.get(v, v)


def parse_index_line(line: str, section: str) -> IndexRow | None:
    # Pattern:
    # S | Nome — Camarilla • Ventrue. Papel. Atua em X. Vínculo mais usado: Y.
    # Apoio | Nome — Camarilla • Mortal. Papel. Atua em X. Vínculo mais usado: Y.
    line = dequote_name(line).strip()
    m = re.match(
        r"^(?P<tier>[^|]+)\|\s*(?P<name>.+?)\s+—\s+(?P<sect>[^•]+)\s+•\s+(?P<clan>[^.]+)\.\s+(?P<role>.+?)\.\s+Atua em\s+(?P<domain>.+?)\.\s+V[ií]nculo mais usado:\s*(?P<link>.+?)\.?\s*$",
        line,
    )
    if not m:
        return None
    return IndexRow(
        tier=m.group("tier").strip(),
        name=m.group("name").strip(),
        sect=norm_sect(m.group("sect")),
        clan=norm_clan(m.group("clan")),
        role=m.group("role").strip(),
        domain=m.group("domain").strip(),
        vínculo=m.group("link").strip(),
        from_section=section,
    )


def extract_essential_30(doc: Document) -> set[str]:
    out: set[str] = set()
    in_84 = False
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if st == "Heading 2" and "8.4" in t:
            in_84 = True
            continue
        if in_84 and st == "Heading 2" and t.startswith("8.5"):
            break
        if in_84 and st == "Heading 3" and t:
            out.add(norm(t))
    return out


def extract_rows_82_83(doc: Document) -> list[IndexRow]:
    rows: list[IndexRow] = []
    in_target = False
    section = ""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        st = p.style.name if p.style else ""
        if st == "Heading 2" and ("8.2" in t or "8.3" in t):
            in_target = True
            section = t
            continue
        if in_target and st == "Heading 2" and ("8.4" in t):
            break
        if not in_target:
            continue
        if st.startswith("List"):
            r = parse_index_line(t, section=section)
            if r:
                rows.append(r)
    return rows


def extract_coteries_64(doc: Document) -> dict[str, list[str]]:
    # Capture "6.4.x — <coterie>" + bullets under "Quem é quem"
    result: dict[str, list[str]] = {}
    in_64 = False
    cur = ""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        st = p.style.name if p.style else ""
        if st == "Heading 2" and "6.4" in t:
            in_64 = True
            continue
        if in_64 and st == "Heading 2" and t.startswith("6.5"):
            break
        if not in_64:
            continue
        if st == "Heading 3" and "6.4." in t:
            cur = dequote_name(t)
            result.setdefault(cur, [])
            continue
        if cur and st.startswith("List"):
            # "Nome (Clã) — papel..."
            m = re.match(r"^([^—(]+)", dequote_name(t))
            if m:
                name = m.group(1).strip()
                if name:
                    result[cur].append(name)
    return result


def coterie_id_for_heading(h: str) -> str | None:
    n = norm(h)
    if "coterie ferrugem" in n:
        return "coterie_ferrugem_mooca"
    if "leste de aco" in n:
        return "coterie_leste_de_aco"
    if "matilha do sul" in n:
        return "coterie_matilha_do_sul"
    return None


def find_entity_by_name(entities: list[dict[str, Any]], name: str) -> dict[str, Any] | None:
    key = norm(name)
    alias_map = {
        norm('Damiao "Guarita"'): norm('Damiao "Portaria"'),
        norm('Wesley "Faixa" Dias'): norm('Wesley "Corredor" Dias'),
    }
    key_alias = alias_map.get(key)
    # exact
    for e in entities:
        if norm(e.get("display_name", "")) == key:
            return e
    if key_alias:
        for e in entities:
            if norm(e.get("display_name", "")) == key_alias:
                return e
    # contains fallback
    for e in entities:
        dn = norm(e.get("display_name", ""))
        if key in dn or dn in key:
            return e
    # first+last token fallback for nickname drift (e.g. "Wesley Faixa Dias" vs "Wesley Corredor Dias")
    toks = [t for t in re.split(r"\s+", key) if t]
    if len(toks) >= 2:
        first = toks[0]
        last = toks[-1]
        for e in entities:
            dn = norm(e.get("display_name", ""))
            d_toks = [t for t in re.split(r"\s+", dn) if t]
            if len(d_toks) >= 2 and d_toks[0] == first and d_toks[-1] == last:
                return e
    return None


def main() -> int:
    doc = Document(DOCX)
    source = json.loads(SRC.read_text(encoding="utf-8"))
    entities: list[dict[str, Any]] = source.get("entities") or []
    coteries: list[dict[str, Any]] = source.get("coteries") or []

    essential = extract_essential_30(doc)
    rows = extract_rows_82_83(doc)
    cot = extract_coteries_64(doc)

    updated = 0
    unresolved_rows: list[str] = []
    link_added = 0
    coterie_assign = 0

    for r in rows:
        e = find_entity_by_name(entities, r.name)
        if not e:
            unresolved_rows.append(r.name)
            continue
        if norm(e.get("display_name", "")) in essential:
            # Keep 8.4 detail import as source of truth for this group.
            continue

        changed_local = False
        if r.tier.strip().upper() in {"S", "A", "B", "C"} and e.get("tier") != r.tier.strip().upper():
            e["tier"] = r.tier.strip().upper()
            changed_local = True
        if r.sect and e.get("sect") != r.sect:
            e["sect"] = r.sect
            changed_local = True
        if r.clan and e.get("clan") != r.clan and e.get("kind") != "ghoul":
            e["clan"] = r.clan
            changed_local = True
        if r.role and e.get("role") != r.role:
            e["role"] = r.role
            changed_local = True
        if r.domain and e.get("domain") != r.domain:
            e["domain"] = r.domain
            changed_local = True

        # Persist the "vinculo mais usado" as a canonical relationship hint.
        target = find_entity_by_name(entities, r.vínculo)
        if target and target.get("id") and e.get("id"):
            links = e.setdefault("links", [])
            exists = any((lk.get("to") == target["id"]) for lk in links if isinstance(lk, dict))
            if not exists:
                links.append(
                    {
                        "to": target["id"],
                        "type": "canon_anchor",
                        "note": f"Vinculo mais usado (DOCX canon, {r.from_section})",
                    }
                )
                link_added += 1
                changed_local = True

        if changed_local:
            updated += 1

    # Apply canonical coterie membership from chapter 6.4
    by_id = {e.get("id"): e for e in entities if e.get("id")}
    for heading, names in cot.items():
        cid = coterie_id_for_heading(heading)
        if not cid:
            continue
        c = next((x for x in coteries if x.get("id") == cid), None)
        if not c:
            continue
        members = c.setdefault("members", [])
        for n in names:
            e = find_entity_by_name(entities, n)
            if not e or not e.get("id"):
                continue
            if e["id"] not in members:
                members.append(e["id"])
                coterie_assign += 1
            # Mirror in entity coteries list.
            ec = e.setdefault("coteries", [])
            if cid not in ec:
                ec.append(cid)

    # Refresh docs snippets for non-essential rows that changed materially.
    for e in entities:
        if norm(e.get("display_name", "")) in essential:
            continue
        docs = e.setdefault("docs", {}).setdefault("files", {})
        # lightweight refresh, keep existing long history if present
        if isinstance(docs.get("ficha_resumida"), str) and docs["ficha_resumida"].strip():
            lines = docs["ficha_resumida"].splitlines()
            out = []
            for ln in lines:
                if ln.startswith("Cla:"):
                    out.append(f"Cla: {e.get('clan','-')} | Seita: {e.get('sect','-')}")
                    continue
                if ln.startswith("Dominio / area principal:"):
                    out.append(f"Dominio / area principal: {e.get('domain','-')}")
                    continue
                out.append(ln)
            docs["ficha_resumida"] = "\n".join(out).rstrip() + "\n"

    source["entities"] = entities
    source["coteries"] = coteries
    source.setdefault("meta", {})
    source["meta"]["canon_index_sync"] = DOCX.name
    source["meta"]["canon_index_sync_at"] = "2026-02-20"

    SRC.write_text(json.dumps(source, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    rep = []
    rep.append("# Relatorio de Cruzamento (DOCX 8.2/8.3 + 6.4)")
    rep.append("")
    rep.append(f"- arquivo: `{DOCX.name}`")
    rep.append(f"- linhas_index_processadas: **{len(rows)}**")
    rep.append(f"- entidades_atualizadas (fora top30): **{updated}**")
    rep.append(f"- links_canon_anchor_adicionados: **{link_added}**")
    rep.append(f"- atribuicoes_de_coterie (cap. 6.4): **{coterie_assign}**")
    rep.append(f"- nao_resolvidos_no_index: **{len(unresolved_rows)}**")
    if unresolved_rows:
        for n in unresolved_rows[:40]:
            rep.append(f"  - {n}")
    REPORT.write_text("\n".join(rep).rstrip() + "\n", encoding="utf-8", newline="\n")

    print(f"[index-canon] rows={len(rows)} updated={updated} unresolved={len(unresolved_rows)}")
    print(f"[index-canon] coterie_assign={coterie_assign} links_added={link_added}")
    print(f"[index-canon] wrote {SRC}")
    print(f"[index-canon] wrote {REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
