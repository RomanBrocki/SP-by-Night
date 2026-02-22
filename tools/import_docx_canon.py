from __future__ import annotations

import json
import os
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = Path(os.environ.get("SPBN_DOCX", str(ROOT / "Sao_Paulo_by_Night.docx")))
SRC = ROOT / "tools" / "sp_by_night_source.json"
REPORT = ROOT / "tools" / "docx_canon_import_report.md"
CANON_MD = ROOT / "01_BACKGROUND_NARRADOR" / "canon_docx_integral.md"


def norm(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace('"', "").replace("'", "")
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    s = re.sub(r"\s+", " ", s)
    return s


def to_md(doc: Document) -> str:
    out: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            out.append("")
            continue
        st = p.style.name if p.style else ""
        if st.startswith("Heading "):
            try:
                lvl = int(st.split()[-1])
            except Exception:
                lvl = 2
            lvl = max(1, min(6, lvl))
            out.append("#" * lvl + " " + t)
            continue
        if st.startswith("List"):
            out.append("- " + t)
            continue
        out.append(t)
    return "\n".join(out).rstrip() + "\n"


@dataclass
class NpcCanon:
    name: str
    tier: str | None = None
    sect: str | None = None
    clan: str | None = None
    role: str | None = None
    kind: str | None = None
    domain: str | None = None
    blood_potency: int | None = None
    humanity: int | None = None
    predator: str | None = None
    disciplines: list[str] = field(default_factory=list)
    apparent_age: str | None = None
    born_year: int | None = None
    embrace_year: int | None = None
    sire: str | None = None
    childer: list[str] = field(default_factory=list)
    appearance: str | None = None
    signature: str | None = None
    ambition: str | None = None
    fear: str | None = None
    secret: str | None = None
    dangerous_truth: str | None = None
    false_rumor: str | None = None
    scene_hooks: list[str] = field(default_factory=list)


def parse_header_line(line: str, npc: NpcCanon) -> None:
    # Example: "S • Camarilla • Ventrue • Príncipe de São Paulo (Kindred)"
    parts = [p.strip() for p in line.split("•")]
    if len(parts) >= 4:
        npc.tier = parts[0]
        npc.sect = parts[1]
        npc.clan = parts[2]
        role_kind = parts[3]
        m = re.match(r"^(.*)\(([^)]+)\)\s*$", role_kind)
        if m:
            npc.role = m.group(1).strip()
            npc.kind = m.group(2).strip().lower()
        else:
            npc.role = role_kind


def parse_domain_line(line: str, npc: NpcCanon) -> None:
    m = re.search(r"Dom[ií]nio:\s*([^\.]+)", line)
    if m:
        npc.domain = m.group(1).strip()


def parse_stats_line(line: str, npc: NpcCanon) -> None:
    m = re.search(r"Pot[êe]ncia de Sangue:\s*(\d+)", line)
    if m:
        npc.blood_potency = int(m.group(1))
    m = re.search(r"Humanidade:\s*(\d+)", line)
    if m:
        npc.humanity = int(m.group(1))
    m = re.search(r"Predador:\s*([^\n]+)$", line)
    if m:
        npc.predator = m.group(1).strip()


def parse_disciplines_line(line: str, npc: NpcCanon) -> None:
    m = re.search(r"Disciplinas principais:\s*([^•\n]+)", line)
    if m:
        npc.disciplines = [x.strip() for x in m.group(1).split(",") if x.strip()]


def parse_history_line(line: str, npc: NpcCanon) -> None:
    m = re.search(r"Nasc\.\s*:\s*(\d{4})", line)
    if m:
        npc.born_year = int(m.group(1))
    m = re.search(r"Abra[çc]o:\s*(\d{4})", line)
    if m:
        npc.embrace_year = int(m.group(1))
    m = re.search(r"Idade aparente:\s*(\d+)", line)
    if m:
        npc.apparent_age = m.group(1)


def parse_sire_line(line: str, npc: NpcCanon) -> None:
    m = re.match(r"Sire:\s*(.+)$", line)
    if m:
        v = m.group(1).strip()
        npc.sire = "" if v in {"—", "-", "N/A"} else v


def parse_childer_line(line: str, npc: NpcCanon) -> None:
    m = re.match(r"Childe\(s\):\s*(.+)$", line)
    if not m:
        return
    v = m.group(1).strip()
    if v in {"—", "-", "N/A"}:
        npc.childer = []
        return
    npc.childer = [x.strip() for x in re.split(r",|;", v) if x.strip()]


def parse_presence_line(line: str, npc: NpcCanon) -> None:
    m = re.search(r"Presen[çc]a em cena:\s*([^;]+)", line)
    if m:
        npc.appearance = m.group(1).strip()
    m = re.search(r"Assinatura social:\s*(.+)$", line)
    if m:
        npc.signature = m.group(1).strip()


def extract_npcs(doc: Document) -> list[NpcCanon]:
    paras = doc.paragraphs
    start = None
    end = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if st == "Heading 2" and "8.4" in t:
            start = i
        elif start is not None and st == "Heading 2" and t.startswith("8.5"):
            end = i
            break
    if start is None:
        raise SystemExit("Nao encontrou secao 8.4 no docx")
    if end is None:
        end = len(paras)

    entries: list[tuple[str, list[tuple[str, str]]]] = []
    cur_name = None
    cur: list[tuple[str, str]] = []
    for p in paras[start:end]:
        t = (p.text or "").strip()
        if not t:
            continue
        st = p.style.name if p.style else ""
        if st == "Heading 3":
            if cur_name:
                entries.append((cur_name, cur))
            cur_name = t
            cur = []
            continue
        if cur_name:
            cur.append((st, t))
    if cur_name:
        entries.append((cur_name, cur))

    out: list[NpcCanon] = []
    for name, lines in entries:
        npc = NpcCanon(name=name)
        mode = ""
        for st, line in lines:
            if "Ganchos curtos" in line:
                mode = "hooks"
                continue
            if st.startswith("List") and mode == "hooks":
                npc.scene_hooks.append(line)
                continue
            if st.startswith("Heading"):
                continue
            parse_header_line(line, npc)
            parse_domain_line(line, npc)
            parse_stats_line(line, npc)
            parse_disciplines_line(line, npc)
            parse_history_line(line, npc)
            parse_sire_line(line, npc)
            parse_childer_line(line, npc)
            parse_presence_line(line, npc)
            if line.startswith("Objetivo"):
                npc.ambition = line.split(":", 1)[1].strip() if ":" in line else line
            elif line.startswith("Medo"):
                npc.fear = line.split(":", 1)[1].strip() if ":" in line else line
            elif line.startswith("Segredo"):
                npc.secret = line.split(":", 1)[1].strip() if ":" in line else line
            elif line.startswith("Verdade perigosa"):
                npc.dangerous_truth = line.split(":", 1)[1].strip() if ":" in line else line
            elif line.startswith("Falso rumor"):
                npc.false_rumor = line.split(":", 1)[1].strip() if ":" in line else line
        out.append(npc)
    return out


def norm_sect(v: str | None) -> str | None:
    if not v:
        return None
    s = norm(v)
    if "camar" in s:
        return "Camarilla"
    if "anar" in s:
        return "Anarch"
    if "autar" in s:
        return "Autarquicos"
    if "indep" in s:
        return "Independentes"
    return v


def norm_kind(v: str | None) -> str | None:
    if not v:
        return None
    s = norm(v)
    if "kindred" in s or "cainit" in s or "vamp" in s:
        return "kindred"
    if "ghoul" in s:
        return "ghoul"
    if "mortal" in s:
        return "mortal"
    return None


def build_docs_text(entity: dict[str, Any], npc: NpcCanon) -> tuple[str, str]:
    display = entity.get("display_name") or npc.name
    alias = entity.get("role") or npc.role or "-"
    clan = entity.get("clan") or npc.clan or "-"
    sect = entity.get("sect") or npc.sect or "-"
    age = entity.get("apparent_age")
    age_line = f"Idade aparente: {age} | " if age else ""
    embrace = entity.get("embrace_year") or "-"
    bp = ((entity.get("full_sheet") or {}).get("blood_potency")) if isinstance(entity.get("full_sheet"), dict) else None
    bp = bp if bp is not None else "-"
    domain = entity.get("domain") or npc.domain or "-"
    discs = entity.get("disciplines_top3") or npc.disciplines or []
    sig = entity.get("signature_style") or npc.signature or "-"
    hooks = entity.get("scene_hooks") or npc.scene_hooks or []
    touchstones = entity.get("touchstones") or []
    sire = entity.get("sire") or npc.sire or "-"
    childer = entity.get("childer") or npc.childer or []
    fear = entity.get("fear") or npc.fear or "-"
    ambition = entity.get("ambition") or npc.ambition or "-"
    secret = entity.get("secret") or npc.secret or "-"
    rumor = entity.get("false_rumor") or npc.false_rumor or "-"
    truth = entity.get("dangerous_truth") or npc.dangerous_truth or "-"

    ficha = []
    ficha.append(f"Nome / Alcunha: {display} / {alias}")
    ficha.append(f"Cla: {clan} | Seita: {sect}")
    ficha.append(f"{age_line}Abraco: {embrace} | Potencia do Sangue {bp}")
    ficha.append(f"Dominio / area principal: {domain}")
    ficha.append(f"Disciplinas (top 3): {discs}")
    ficha.append(f"Assinatura / estilo: {sig}")
    if touchstones:
        ficha.append("Touchstones/Conviccoes (resumo):")
        for t in touchstones[:2]:
            ficha.append(f"- {t}")
    if hooks:
        ficha.append("Ganchos de cena:")
        for h in hooks[:3]:
            ficha.append(f"- {h}")
    ficha_txt = "\n".join(ficha).rstrip() + "\n"

    hist = []
    hist.append(display)
    hist.append("")
    hist.append("Perfil:")
    hist.append(f"- Cla: {clan} | Seita: {sect} | Funcao: {alias}")
    dom_age = f"- Dominio: {domain}"
    if age:
        dom_age += f" | Idade aparente: {age}"
    hist.append(dom_age)
    hist.append(f"- Abraco: {embrace} | Sire: {sire} | Acima do sire: -")
    hist.append(f"- Disciplinas (top 3): {discs}")
    hist.append(f"- Assinatura/estilo: {sig}")
    hist.append("")
    hist.append("Ambicao, Medo e Segredo:")
    hist.append(f"- Ambicao: {ambition}")
    hist.append(f"- Medo: {fear}")
    hist.append(f"- Segredo: {secret}")
    hist.append("")
    hist.append("Rumor e verdade perigosa:")
    hist.append(f"- Rumor falso: {rumor}")
    hist.append(f"- Verdade perigosa: {truth}")
    if childer:
        hist.append("")
        hist.append("Descendencia:")
        for c in childer:
            hist.append(f"- Childe: {c}")
    hist_txt = "\n".join(hist).rstrip() + "\n"
    return ficha_txt, hist_txt


def main() -> int:
    if not DOCX.exists():
        raise SystemExit(f"Arquivo nao encontrado: {DOCX}")
    if not SRC.exists():
        raise SystemExit(f"Arquivo nao encontrado: {SRC}")

    doc = Document(DOCX)
    CANON_MD.write_text(to_md(doc), encoding="utf-8", newline="\n")

    source = json.loads(SRC.read_text(encoding="utf-8"))
    entities: list[dict[str, Any]] = source.get("entities") or []
    by_name = {norm(e.get("display_name", "")): e for e in entities}

    npcs = extract_npcs(doc)

    changed_fields = 0
    changed_entities = 0
    unmatched: list[str] = []
    lines = []
    lines.append("# Relatorio de Importacao do DOCX Canon")
    lines.append("")
    lines.append(f"- arquivo: `{DOCX.name}`")
    lines.append(f"- npc_essenciais_detectados: **{len(npcs)}**")
    lines.append("")

    for npc in npcs:
        key = norm(npc.name)
        e = by_name.get(key)
        if not e:
            unmatched.append(npc.name)
            continue

        before = json.dumps(
            {
                "role": e.get("role"),
                "domain": e.get("domain"),
                "sect": e.get("sect"),
                "tier": e.get("tier"),
                "ambition": e.get("ambition"),
                "fear": e.get("fear"),
                "secret": e.get("secret"),
                "dangerous_truth": e.get("dangerous_truth"),
                "false_rumor": e.get("false_rumor"),
                "scene_hooks": e.get("scene_hooks"),
            },
            ensure_ascii=False,
            sort_keys=True,
        )

        if npc.role:
            e["role"] = npc.role
        if npc.domain:
            e["domain"] = npc.domain
        if npc.sect:
            e["sect"] = norm_sect(npc.sect) or npc.sect
        if npc.clan:
            e["clan"] = npc.clan.replace("ã", "a").replace("ç", "c") if e.get("clan") == "BrÃ¼jah" else npc.clan
        if npc.tier:
            e["tier"] = npc.tier
        k = norm_kind(npc.kind)
        if k:
            e["kind"] = k
        if npc.ambition:
            e["ambition"] = npc.ambition
        if npc.fear:
            e["fear"] = npc.fear
        if npc.secret:
            e["secret"] = npc.secret
        if npc.dangerous_truth:
            e["dangerous_truth"] = npc.dangerous_truth
        if npc.false_rumor:
            e["false_rumor"] = npc.false_rumor
        if npc.scene_hooks:
            e["scene_hooks"] = npc.scene_hooks[:3]
        if npc.disciplines:
            e["disciplines_top3"] = npc.disciplines[:3]
        if npc.signature:
            e["signature_style"] = npc.signature
        if npc.appearance:
            e["appearance_explicit"] = npc.appearance
        if npc.sire is not None:
            e["sire"] = npc.sire
        if npc.childer:
            e["childer"] = npc.childer
        if npc.born_year:
            e["born_year"] = npc.born_year
        if npc.embrace_year:
            e["embrace_year"] = npc.embrace_year
        if npc.apparent_age and (e.get("kind") != "kindred" or norm(e.get("clan", "")) != "nosferatu"):
            # keep current project rule: max 40
            e["apparent_age"] = str(min(int(npc.apparent_age), 40))
        if e.get("kind") == "kindred" and norm(e.get("clan", "")) == "nosferatu":
            e.pop("apparent_age", None)

        fs = e.get("full_sheet")
        if isinstance(fs, dict):
            if npc.blood_potency is not None:
                fs["blood_potency"] = npc.blood_potency
            if npc.humanity is not None:
                fs["humanity"] = npc.humanity
            if npc.predator:
                fs["predator_type"] = npc.predator

        docs = e.setdefault("docs", {}).setdefault("files", {})
        ficha, hist = build_docs_text(e, npc)
        docs["ficha_resumida"] = ficha
        docs["historia"] = hist

        after = json.dumps(
            {
                "role": e.get("role"),
                "domain": e.get("domain"),
                "sect": e.get("sect"),
                "tier": e.get("tier"),
                "ambition": e.get("ambition"),
                "fear": e.get("fear"),
                "secret": e.get("secret"),
                "dangerous_truth": e.get("dangerous_truth"),
                "false_rumor": e.get("false_rumor"),
                "scene_hooks": e.get("scene_hooks"),
            },
            ensure_ascii=False,
            sort_keys=True,
        )
        if before != after:
            changed_entities += 1
            changed_fields += sum(1 for a, b in zip(before, after) if a != b)
            lines.append(f"- atualizado: **{e.get('display_name')}** (`{e.get('id')}`)")

    source["meta"] = source.get("meta") or {}
    source["meta"]["canon_source"] = DOCX.name
    source["meta"]["canon_imported_at"] = "2026-02-20"

    SRC.write_text(json.dumps(source, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    lines.append("")
    lines.append(f"- entidades_com_mudancas: **{changed_entities}**")
    lines.append(f"- campos_alterados_aprox: **{changed_fields}**")
    if unmatched:
        lines.append(f"- nao_encontrados: **{len(unmatched)}**")
        for n in unmatched:
            lines.append(f"  - {n}")
    else:
        lines.append("- nao_encontrados: **0**")
    REPORT.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8", newline="\n")

    print(f"[canon] npcs={len(npcs)} changed_entities={changed_entities} unmatched={len(unmatched)}")
    print(f"[canon] wrote {SRC}")
    print(f"[canon] wrote {REPORT}")
    print(f"[canon] wrote {CANON_MD}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
