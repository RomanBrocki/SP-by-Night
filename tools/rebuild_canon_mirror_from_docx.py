from __future__ import annotations

import json
import re
import unicodedata
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Sao_Paulo_by_Night.docx"
MIRROR_ROOT = ROOT / "RECONSTRUCAO_CANONICA"
PROMPT_SOURCE_DIR = ROOT / "05_ASSETS" / "portrait_prompts"
PROMPT_MIRROR_DIR = MIRROR_ROOT / "05_ASSETS" / "portrait_prompts"


def collapse_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def norm_key(text: str, drop_articles: bool = True) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    plain = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    plain = (
        plain.replace("–", " ")
        .replace("—", " ")
        .replace("-", " ")
        .replace("/", " ")
        .replace('"', " ")
        .replace("'", " ")
    )
    tokens = re.findall(r"[a-z0-9]+", plain.lower())
    if drop_articles:
        tokens = [token for token in tokens if token not in {"o", "a"}]
    return " ".join(tokens)


LEGACY_FILE_ALIASES = {
    norm_key("Damião Guarita"): "Damiao_Portaria",
    norm_key('Wesley "Faixa" Dias'): "Wesley_Corredor_Dias",
}


@dataclass
class ParagraphRec:
    index: int
    style: str
    text: str


@dataclass
class HeadingNode:
    title: str
    level: int
    index: int
    parent: "HeadingNode | None" = None
    paragraphs: list[ParagraphRec] = field(default_factory=list)
    children: list["HeadingNode"] = field(default_factory=list)


@dataclass
class IndexEntry:
    source: str
    section: str
    subgroup: str | None
    raw: str
    tier: str | None
    name: str
    faction: str | None
    classification: str | None
    role: str | None
    domain: str | None
    link: str | None
    parse_failed: bool = False


@dataclass
class EssentialEntry:
    name: str
    node: HeadingNode

    def lines(self) -> list[ParagraphRec]:
        return self.node.paragraphs

    def first_line(self) -> str | None:
        for paragraph in self.lines():
            if not paragraph.style.startswith("List Bullet"):
                return paragraph.text
        return None

    def find_line(self, prefix: str) -> str | None:
        prefix_key = norm_key(prefix, drop_articles=False)
        for paragraph in self.lines():
            head = paragraph.text.split(":", 1)[0]
            if norm_key(head, drop_articles=False) == prefix_key:
                return paragraph.text
            if norm_key(paragraph.text, drop_articles=False).startswith(prefix_key):
                return paragraph.text
        return None

    def bullets_after(self, header: str) -> list[str]:
        items: list[str] = []
        current: str | None = None
        header_key = norm_key(header, drop_articles=False)
        for paragraph in self.lines():
            if paragraph.style.startswith("List Bullet"):
                if current == header_key:
                    items.append(paragraph.text)
                continue
            current = None
            if paragraph.text.endswith(":") and ":" not in paragraph.text[:-1]:
                current = norm_key(paragraph.text[:-1], drop_articles=False)
        return items


def heading_level(style_name: str) -> int | None:
    match = re.fullmatch(r"Heading ([1-6])", style_name)
    if not match:
        return None
    return int(match.group(1))


def build_tree(doc: Document) -> HeadingNode:
    root = HeadingNode(title="ROOT", level=0, index=-1, parent=None)
    stack = [root]
    for index, paragraph in enumerate(doc.paragraphs):
        text = collapse_ws(paragraph.text)
        if not text:
            continue
        level = heading_level(paragraph.style.name)
        if level is not None:
            while stack and stack[-1].level >= level:
                stack.pop()
            node = HeadingNode(title=text, level=level, index=index, parent=stack[-1])
            stack[-1].children.append(node)
            stack.append(node)
            continue
        stack[-1].paragraphs.append(
            ParagraphRec(index=index, style=paragraph.style.name, text=text)
        )
    return root


def find_child_by_prefix(node: HeadingNode, prefix: str) -> HeadingNode:
    target = norm_key(prefix, drop_articles=False)
    for child in node.children:
        if norm_key(child.title, drop_articles=False).startswith(target):
            return child
    raise KeyError(f"Heading not found below {node.title!r}: {prefix!r}")


def find_path(root: HeadingNode, prefixes: list[str]) -> HeadingNode:
    current = root
    for prefix in prefixes:
        current = find_child_by_prefix(current, prefix)
    return current


def format_paragraph(paragraph: ParagraphRec) -> str:
    if paragraph.style.startswith("List Bullet"):
        return f"- {paragraph.text}"
    return paragraph.text


def tidy_lines(lines: list[str]) -> str:
    cleaned: list[str] = []
    previous_blank = True
    for line in lines:
        text = line.rstrip()
        if not text:
            if previous_blank:
                continue
            cleaned.append("")
            previous_blank = True
            continue
        cleaned.append(text)
        previous_blank = False
    while cleaned and not cleaned[-1]:
        cleaned.pop()
    return "\n".join(cleaned) + "\n"


def render_node(node: HeadingNode, base_level: int | None = None) -> list[str]:
    if base_level is None:
        base_level = node.level
    lines = ["#" * (node.level - base_level + 1) + f" {node.title}", ""]
    for paragraph in node.paragraphs:
        lines.append(format_paragraph(paragraph))
        lines.append("")
    for child in node.children:
        lines.extend(render_node(child, base_level))
    return lines


def render_compilation(title: str, source_note: str, nodes: list[HeadingNode]) -> str:
    lines = [f"# {title}", "", f"Fonte canônica: {source_note}", ""]
    for index, node in enumerate(nodes):
        if index:
            lines.append("")
        lines.extend(render_node(node))
    return tidy_lines(lines)


ENTRY_HEAD_RE = re.compile(
    r"^(?P<tier>[^|]+?)\s*\|\s*"
    r"(?P<name>.+?)\s+-\s+"
    r"(?P<faction>[^•]+?)\s*•\s*"
    r"(?P<classification>.+?)\.\s*"
    r"(?P<rest>.+?)\.?$"
)


def parse_index_line(
    text: str,
    source: str,
    section: str,
    subgroup: str | None,
) -> IndexEntry:
    match = ENTRY_HEAD_RE.match(text)
    if not match:
        return IndexEntry(
            source=source,
            section=section,
            subgroup=subgroup,
            raw=text,
            tier=None,
            name=text,
            faction=None,
            classification=None,
            role=None,
            domain=None,
            link=None,
            parse_failed=True,
        )

    rest = match.group("rest").strip()
    sentences = [part.strip() for part in re.split(r"\.\s+", rest) if part.strip()]
    role = sentences[0] if sentences else None
    domain = None
    link = None
    for sentence in sentences[1:]:
        if sentence.startswith("Atua "):
            domain = re.sub(r"^Atua\s+(?:em|na|no)\s+", "", sentence).strip()
            continue
        if sentence.startswith("Vínculo mais usado:"):
            link = sentence.split(":", 1)[1].strip()

    return IndexEntry(
        source=source,
        section=section,
        subgroup=subgroup,
        raw=text,
        tier=match.group("tier"),
        name=match.group("name"),
        faction=match.group("faction").strip(),
        classification=match.group("classification").strip(),
        role=role,
        domain=domain,
        link=link,
        parse_failed=role is None or link is None,
    )


def collect_index_entries(section_node: HeadingNode, source: str) -> list[IndexEntry]:
    entries: list[IndexEntry] = []
    for child in section_node.children:
        for paragraph in child.paragraphs:
            if paragraph.style.startswith("List Bullet"):
                entries.append(parse_index_line(paragraph.text, source, child.title, None))
        for grandchild in child.children:
            for paragraph in grandchild.paragraphs:
                if paragraph.style.startswith("List Bullet"):
                    entries.append(
                        parse_index_line(
                            paragraph.text,
                            source,
                            child.title,
                            grandchild.title,
                        )
                    )
    return entries


def collect_essentials(section_node: HeadingNode) -> dict[str, EssentialEntry]:
    return {
        child.title: EssentialEntry(name=child.title, node=child)
        for child in section_node.children
    }


def quote_yaml(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def write_text(path: Path, text: str) -> None:
    ensure_parent(path)
    path.write_text(text, encoding="utf-8")


def ensure_prompt_mirror_files() -> list[Path]:
    created: list[Path] = []
    if not PROMPT_SOURCE_DIR.exists():
        return created
    for source_path in sorted(PROMPT_SOURCE_DIR.glob("*.txt")):
        target_path = PROMPT_MIRROR_DIR / source_path.name
        if not target_path.exists():
            ensure_parent(target_path)
            target_path.write_text("", encoding="utf-8")
            created.append(target_path)
    return created


def build_lookup(paths: list[Path], suffix: str | None = None) -> dict[str, list[Path]]:
    lookup: dict[str, list[Path]] = defaultdict(list)
    for path in paths:
        base = path.stem
        if suffix and base.endswith(suffix):
            base = base[: -len(suffix)]
        lookup[norm_key(base)].append(path)
    return lookup


def resolve_paths(name: str, lookup: dict[str, list[Path]]) -> list[Path]:
    key = norm_key(name)
    if key in lookup:
        return sorted(lookup[key])
    alias = LEGACY_FILE_ALIASES.get(key)
    if alias:
        return sorted(lookup.get(norm_key(alias), []))
    return []


def format_index_summary(entry: IndexEntry) -> str:
    if entry.parse_failed:
        return tidy_lines(
            [
                f"Fonte canônica: {entry.source}",
                "",
                "Linha do livro:",
                entry.raw,
            ]
        )
    lines = [
        f"Fonte canônica: {entry.source}",
        "",
        f"Nome: {entry.name}",
        f"Tier: {entry.tier}",
        f"Seção maior: {entry.section}",
    ]
    if entry.subgroup:
        lines.append(f"Subgrupo no livro: {entry.subgroup}")
    lines.extend(
        [
            f"Facção/afiliação: {entry.faction}",
            f"Clã ou tipo: {entry.classification}",
            f"Função: {entry.role}",
            f"Área de atuação: {entry.domain}",
            f"Vínculo mais usado: {entry.link}",
            "",
            "Linha literal do índice:",
            entry.raw,
        ]
    )
    return tidy_lines(lines)


def format_index_history(entry: IndexEntry) -> str:
    if entry.parse_failed:
        return tidy_lines(
            [
                f"Fonte canônica: {entry.source}",
                "",
                "O livro, nesta etapa, só oferece a linha abaixo para este personagem:",
                entry.raw,
            ]
        )
    return tidy_lines(
        [
            f"Fonte canônica: {entry.source}",
            "",
            "O livro, nesta etapa, não traz ficha expandida própria para este personagem no material extraído.",
            "",
            f"Nome: {entry.name}",
            f"Linha de índice: {entry.raw}",
        ]
    )


def format_essential_summary(
    essential: EssentialEntry,
    index_entry: IndexEntry | None,
) -> str:
    lines = ["Fonte canônica: 8.4 - As fichas essenciais"]
    if index_entry:
        lines.append("Apoio complementar: 8.2/8.3 - índices canônicos")
    lines.append("")

    first = essential.first_line()
    if first:
        lines.append(first)
    for prefix in (
        "Domínio:",
        "Estado:",
        "Potância de Sangue:",
        "Disciplinas principais:",
        "Disciplinas principais (ghoul):",
        "Paradas de dados uteis:",
        "Nasc.:",
        "Sire:",
        "Childe(s):",
        "Domitor/Patrocinador:",
        "Equipe e dependentes:",
        "Objetivo:",
        "Medo:",
        "Segredo:",
        "Verdade perigosa:",
        "Falso rumor que cola:",
    ):
        value = essential.find_line(prefix)
        if value:
            lines.append(value)

    hooks = essential.bullets_after("Ganchos curtos")
    if hooks:
        lines.extend(["", "Ganchos curtos:"])
        lines.extend(f"- {hook}" for hook in hooks)

    if index_entry and index_entry.link:
        lines.extend(["", f"Vínculo mais usado no índice: {index_entry.link}"])

    return tidy_lines(lines)


def format_essential_history(
    essential: EssentialEntry,
    index_entry: IndexEntry | None,
) -> str:
    lines = ["Fonte canônica: 8.4 - As fichas essenciais"]
    if index_entry:
        lines.append("Apoio complementar: 8.2/8.3 - índices canônicos")
    lines.extend(["", f"Nome: {essential.name}", ""])

    include = False
    for paragraph in essential.lines():
        if paragraph.text == "História e comportamento:":
            include = True
        if include:
            lines.append(format_paragraph(paragraph))
    return tidy_lines(lines)


def format_essential_full(essential: EssentialEntry) -> str:
    lines = [f"# {essential.name}", "", "Fonte canônica: 8.4 - As fichas essenciais", ""]
    for paragraph in essential.lines():
        lines.append(format_paragraph(paragraph))
        lines.append("")
    return tidy_lines(lines)


def format_prompt_from_essential(
    essential: EssentialEntry,
    index_entry: IndexEntry | None,
) -> str:
    presence = essential.find_line("Presença em cena:")
    if not presence:
        return ""
    lines = [
        "Fonte canônica: 8.4 - As fichas essenciais",
        "",
        f"Nome: {essential.name}",
    ]
    first = essential.first_line()
    if first:
        lines.append(f"Classificação: {first}")
    domain = essential.find_line("Domínio:")
    if domain:
        lines.append(domain)
    lines.append(presence)
    if index_entry and index_entry.link:
        lines.append(f"Vínculo mais usado no índice: {index_entry.link}")
    return tidy_lines(lines)


def slug_folder_label(folder_name: str) -> str:
    return folder_name.replace("_", " ")


def build_clan_structure(clan_label: str, entries: list[IndexEntry]) -> str:
    lines = [
        f"# Estrutura do clã: {clan_label}",
        "",
        "Fonte canônica: 8.2 - Índice completo por facção e clã",
        "",
    ]
    grouped: dict[str, list[IndexEntry]] = defaultdict(list)
    for entry in entries:
        grouped[entry.section].append(entry)
    for section_name in sorted(grouped):
        lines.append(f"## {section_name}")
        lines.append("")
        for entry in sorted(grouped[section_name], key=lambda item: item.name):
            lines.append(f"- {entry.raw}")
        lines.append("")
    return tidy_lines(lines)


def build_faction_index() -> str:
    return tidy_lines(
        [
            "# Índice de Facções",
            "",
            "Fonte canônica principal: capítulos 2, 3, 4, 7 e capítulo 8 do livro.",
            "",
            "- Camarilla -> `faccoes/Camarilla.md`",
            "- Anarquistas -> `faccoes/Anarquistas.md`",
            "- Independentes -> `faccoes/Independentes.md`",
            "- Segunda Inquisição -> `faccoes/Segunda_Inquisicao.md`",
            "- Mortal -> `faccoes/Mortal.md`",
            "",
            "Arquivos espelho com rótulo legado sem seção canônica direta no livro:",
            "- `faccoes/Anarch.md`",
            "- `faccoes/Autarquicos.md`",
        ]
    )


def build_mortal_faction_doc(mortal_entries: list[IndexEntry]) -> str:
    lines = ["# Mortal", "", "Fonte canônica: 8.2 - Índice completo por facção e clã", ""]
    for entry in sorted(mortal_entries, key=lambda item: item.name):
        lines.append(f"- {entry.raw}")
    return tidy_lines(lines)


def build_secrets_doc(essentials: dict[str, EssentialEntry]) -> str:
    lines = [
        "# Segredos e Verdades",
        "",
        "Fonte canônica: 8.4 - As fichas essenciais",
        "",
        "## Segredos",
        "",
    ]
    for name in sorted(essentials):
        line = essentials[name].find_line("Segredo:")
        if line:
            lines.append(f"- {name}: {line.split(':', 1)[1].strip()}")
    lines.extend(["", "## Verdades Perigosas", ""])
    for name in sorted(essentials):
        line = essentials[name].find_line("Verdade perigosa:")
        if line:
            lines.append(f"- {name}: {line.split(':', 1)[1].strip()}")
    return tidy_lines(lines)


def build_rumors_doc(root: HeadingNode, essentials: dict[str, EssentialEntry]) -> str:
    lines = ["# Rumores", "", "Fonte canônica: livro base (capítulos 3, 4 e 8.4)", ""]
    for name in sorted(essentials):
        line = essentials[name].find_line("Falso rumor que cola:")
        if line:
            lines.append(f"- {name}: {line.split(':', 1)[1].strip()}")

    chapter3 = find_path(root, ["Capítulo 3"])
    chapter4 = find_path(root, ["Capítulo 4"])
    for node in chapter3.children + chapter4.children:
        for child in node.children:
            if "boatos" in norm_key(child.title, drop_articles=False):
                lines.extend(["", f"## {child.title}", ""])
                for paragraph in child.paragraphs:
                    lines.append(format_paragraph(paragraph))
    return tidy_lines(lines)


def build_teia_doc(
    entries: list[IndexEntry],
    ghouls: list[IndexEntry],
    essentials: dict[str, EssentialEntry],
) -> str:
    lines = [
        "# Teia de Conexões",
        "",
        "Fonte canônica: 8.2, 8.3 e 8.4",
        "",
        "## Vínculos mais usados",
        "",
    ]
    for entry in sorted(entries + ghouls, key=lambda item: item.name):
        if entry.link:
            lines.append(f"- {entry.name} -> {entry.link}")

    lines.extend(["", "## Linhagens, patrocínios e equipes", ""])
    for name in sorted(essentials):
        essential = essentials[name]
        for prefix in ("Sire:", "Childe(s):", "Domitor/Patrocinador:", "Equipe e dependentes:"):
            line = essential.find_line(prefix)
            if line:
                lines.append(f"- {name}: {line}")
    return tidy_lines(lines)


def build_dominios_yml(
    entries: list[IndexEntry],
    ghouls: list[IndexEntry],
    essentials: dict[str, EssentialEntry],
) -> str:
    lines = ["personagens:"]
    all_entries = {entry.name: entry for entry in entries + ghouls}
    for name in sorted(set(all_entries) | set(essentials)):
        lines.append(f"  {quote_yaml(name)}:")
        source_labels: list[str] = []
        if name in all_entries:
            source_labels.append(all_entries[name].source)
        if name in essentials:
            source_labels.append("8.4 - As fichas essenciais")
        domain_line = essentials[name].find_line("Domínio:") if name in essentials else None
        domain_value = None
        if domain_line and ":" in domain_line:
            domain_value = domain_line.split(":", 1)[1].strip()
        elif name in all_entries:
            domain_value = all_entries[name].domain
        lines.append("    fonte:")
        for label in sorted(dict.fromkeys(source_labels)):
            lines.append(f"      - {quote_yaml(label)}")
        if domain_value:
            lines.append(f"    dominio: {quote_yaml(domain_value)}")
        if name in all_entries and all_entries[name].section:
            lines.append(f"    secao_maior: {quote_yaml(all_entries[name].section)}")
        if name in all_entries and all_entries[name].classification:
            lines.append(f"    cla_ou_tipo: {quote_yaml(all_entries[name].classification)}")
    return "\n".join(lines) + "\n"


def build_rede_yml(
    entries: list[IndexEntry],
    ghouls: list[IndexEntry],
    essentials: dict[str, EssentialEntry],
) -> str:
    by_name = {entry.name: entry for entry in entries + ghouls}
    lines = ["personagens:"]
    for name in sorted(set(by_name) | set(essentials)):
        entry = by_name.get(name)
        essential = essentials.get(name)
        lines.append(f"  {quote_yaml(name)}:")
        if entry and entry.faction:
            lines.append(f"    seita_ou_afinidade: {quote_yaml(entry.faction)}")
        if entry and entry.classification:
            lines.append(f"    cla_ou_tipo: {quote_yaml(entry.classification)}")
        lines.append("    relacoes_explicitas:")
        if entry and entry.link:
            lines.append(f"      - {quote_yaml(f'Vínculo mais usado: {entry.link}')}")
        if essential:
            for prefix in (
                "Sire:",
                "Childe(s):",
                "Domitor/Patrocinador:",
                "Equipe e dependentes:",
                "Segredo:",
                "Verdade perigosa:",
            ):
                line = essential.find_line(prefix)
                if line:
                    lines.append(f"      - {quote_yaml(line)}")
    return "\n".join(lines) + "\n"


def build_boons_yml(entries: list[IndexEntry], essentials: dict[str, EssentialEntry]) -> str:
    keywords = ("dívida", "divida", "boon", "favor")
    lines = [
        "observacoes:",
        '  - "Arquivo parcial: só registra menções literais do livro a dívida, boon ou favor."',
        "registros:",
    ]
    seen = False
    for entry in sorted(entries, key=lambda item: item.name):
        if any(keyword in norm_key(entry.raw, drop_articles=False) for keyword in keywords):
            seen = True
            lines.append(f"  - personagem: {quote_yaml(entry.name)}")
            lines.append(f"    fonte: {quote_yaml(entry.source)}")
            lines.append(f"    texto: {quote_yaml(entry.raw)}")
    for name in sorted(essentials):
        for paragraph in essentials[name].lines():
            if any(keyword in norm_key(paragraph.text, drop_articles=False) for keyword in keywords):
                seen = True
                lines.append(f"  - personagem: {quote_yaml(name)}")
                lines.append('    fonte: "8.4 - As fichas essenciais"')
                lines.append(f"    texto: {quote_yaml(paragraph.text)}")
    if not seen:
        lines.append('  - personagem: "N/A"')
        lines.append('    fonte: "Livro base"')
        lines.append('    texto: "Nenhuma menção literal rastreada nesta passada."')
    return "\n".join(lines) + "\n"


def build_population_report(
    total_empty: int,
    entries: list[IndexEntry],
    ghouls: list[IndexEntry],
    essentials: dict[str, EssentialEntry],
    parse_failures: list[IndexEntry],
    missing_summary: list[str],
    missing_history: list[str],
    missing_full: list[str],
    missing_prompts: list[str],
    unmatched_files: list[str],
    legacy_name_hits: list[str],
) -> str:
    faction_counts = Counter(entry.section for entry in entries)
    clan_counts = Counter(entry.classification for entry in entries if entry.classification)
    lines = [
        "# Relatório de População e Lacunas",
        "",
        "Fonte canônica: capítulo 8 do livro",
        "",
        f"- NPCs indexados em 8.2: {len(entries)}",
        f"- Ghouls/mortais indexados em 8.3: {len(ghouls)}",
        f"- Fichas essenciais em 8.4: {len(essentials)}",
        f"- Arquivos ainda vazios no espelho: {total_empty}",
        "",
        "## Contagem por grande seção",
        "",
    ]
    for label, count in sorted(faction_counts.items()):
        lines.append(f"- {label}: {count}")
    lines.extend(["", "## Contagem por clã/tipo", ""])
    for label, count in sorted(clan_counts.items()):
        lines.append(f"- {label}: {count}")
    if parse_failures:
        lines.extend(["", "## Linhas de índice que ainda precisam de ajuste de parser", ""])
        for entry in parse_failures:
            lines.append(f"- {entry.source}: {entry.raw}")
    if missing_summary:
        lines.extend(["", "## Personagens canônicos sem arquivo espelho de ficha resumida", ""])
        lines.extend(f"- {name}" for name in missing_summary)
    if missing_history:
        lines.extend(["", "## Personagens canônicos sem arquivo espelho de história", ""])
        lines.extend(f"- {name}" for name in missing_history)
    if missing_full:
        lines.extend(["", "## Fichas essenciais sem arquivo espelho de ficha completa", ""])
        lines.extend(f"- {name}" for name in missing_full)
    if missing_prompts:
        lines.extend(["", "## Personagens essenciais sem prompt espelho correspondente", ""])
        lines.extend(f"- {name}" for name in missing_prompts)
    if unmatched_files:
        lines.extend(["", "## Arquivos espelho sem base direta localizada no livro nesta passada", ""])
        lines.extend(f"- {item}" for item in unmatched_files)
    if legacy_name_hits:
        lines.extend(["", "## Arquivos com nome legado mapeados para nome canônico", ""])
        lines.extend(f"- {item}" for item in legacy_name_hits)
    return tidy_lines(lines)


def build_gap_report(
    missing_summary: list[str],
    missing_history: list[str],
    missing_full: list[str],
    missing_prompts: list[str],
    unmatched_files: list[str],
    legacy_name_hits: list[str],
) -> str:
    lines = [
        "# Lacunas para Decisão",
        "",
        "Este arquivo só registra ausência de suporte canônico direto no livro ou ausência de arquivo espelho correspondente.",
    ]
    if missing_summary:
        lines.extend(["", "## Personagens do livro sem ficha resumida espelho", ""])
        lines.extend(f"- {name}" for name in missing_summary)
    if missing_history:
        lines.extend(["", "## Personagens do livro sem arquivo de história espelho", ""])
        lines.extend(f"- {name}" for name in missing_history)
    if missing_full:
        lines.extend(["", "## Fichas essenciais sem arquivo `*_ficha_completa.md` no espelho", ""])
        lines.extend(f"- {name}" for name in missing_full)
    if missing_prompts:
        lines.extend(["", "## Fichas essenciais sem arquivo de prompt espelho", ""])
        lines.extend(f"- {name}" for name in missing_prompts)
    if unmatched_files:
        lines.extend(["", "## Arquivos espelho ainda vazios por falta de base direta nesta passada", ""])
        lines.extend(f"- {item}" for item in unmatched_files)
    if legacy_name_hits:
        lines.extend(["", "## Arquivos com nome legado já mapeados para personagem canônico", ""])
        lines.extend(f"- {item}" for item in legacy_name_hits)
    return tidy_lines(lines)


def main() -> None:
    ensure_prompt_mirror_files()

    doc = Document(str(DOCX_PATH))
    root = build_tree(doc)

    chapter1 = find_path(root, ["Capítulo 1"])
    chapter2 = find_path(root, ["Capítulo 2"])
    chapter3 = find_path(root, ["Capítulo 3"])
    chapter4 = find_path(root, ["Capítulo 4"])
    chapter5 = find_path(root, ["Capítulo 5"])
    chapter6 = find_path(root, ["Capítulo 6"])
    chapter7 = find_path(root, ["Capítulo 7"])
    chapter8 = find_path(root, ["Capítulo 8"])

    sec_13 = find_path(root, ["Capítulo 1", "1.3"])
    sec_14 = find_path(root, ["Capítulo 1", "1.4"])
    sec_16 = find_path(root, ["Capítulo 1", "1.6"])
    sec_17 = find_path(root, ["Capítulo 1", "1.7"])
    sec_51 = find_path(root, ["Capítulo 5", "5.1"])
    sec_52 = find_path(root, ["Capítulo 5", "5.2"])
    sec_53 = find_path(root, ["Capítulo 5", "5.3"])
    sec_72 = find_path(root, ["Capítulo 7", "7.2"])
    sec_73 = find_path(root, ["Capítulo 7", "7.3"])
    sec_74 = find_path(root, ["Capítulo 7", "7.4"])
    sec_75 = find_path(root, ["Capítulo 7", "7.5"])
    sec_76 = find_path(root, ["Capítulo 7", "7.6"])
    sec_77 = find_path(root, ["Capítulo 7", "7.7"])
    sec_82 = find_path(root, ["Capítulo 8", "8.2"])
    sec_83 = find_path(root, ["Capítulo 8", "8.3"])
    sec_84 = find_path(root, ["Capítulo 8", "8.4"])

    general_docs: dict[str, str] = {
        "00_BACKGROUND_JOGADORES/overview_sao_paulo.md": render_compilation(
            "Overview São Paulo",
            "Capítulo 1 - São Paulo by Night",
            [chapter1],
        ),
        "00_BACKGROUND_JOGADORES/bairros_e_dominios.md": render_compilation(
            "Bairros e Domínios",
            "Capítulo 5 - A Geografia da Noite (seções 5.1, 5.2 e 5.3)",
            [sec_51, sec_52, sec_53],
        ),
        "00_BACKGROUND_JOGADORES/mascarade_e_ameacas.md": render_compilation(
            "Máscara e Ameaças",
            "Capítulo 1 (1.3 e 1.4) + Capítulo 7",
            [sec_13, sec_14, chapter7],
        ),
        "00_BACKGROUND_JOGADORES/seitas_e_politica.md": render_compilation(
            "Seitas e Política",
            "Capítulos 2, 3 e 4",
            [chapter2, chapter3, chapter4],
        ),
        "01_BACKGROUND_NARRADOR/cronologia.md": render_compilation(
            "Cronologia",
            "Capítulo 1 - seção 1.6",
            [sec_16],
        ),
        "01_BACKGROUND_NARRADOR/tramas_em_andamento.md": render_compilation(
            "Tramas em Andamento",
            "Capítulo 1 - seção 1.7",
            [sec_17],
        ),
        "01_BACKGROUND_NARRADOR/coteries_e_associacoes.md": render_compilation(
            "Coteries e Associações",
            "Capítulo 6 - As Coteries",
            [chapter6],
        ),
        "01_BACKGROUND_NARRADOR/geopolitica_territorial.md": render_compilation(
            "Geopolítica Territorial",
            "Capítulo 5 - A Geografia da Noite",
            [chapter5],
        ),
        "01_BACKGROUND_NARRADOR/painel_consolidado_faccoes_e_grupos.md": render_compilation(
            "Painel Consolidado de Facções e Grupos",
            "Capítulo 8 - índices 8.2 e 8.3",
            [sec_82, sec_83],
        ),
        "01_BACKGROUND_NARRADOR/faccoes/Camarilla.md": render_compilation(
            "Camarilla",
            "Capítulo 2 - A Camarilla",
            [chapter2],
        ),
        "01_BACKGROUND_NARRADOR/faccoes/Anarquistas.md": render_compilation(
            "Anarquistas",
            "Capítulo 3 - Os Baronatos",
            [chapter3],
        ),
        "01_BACKGROUND_NARRADOR/faccoes/Independentes.md": render_compilation(
            "Independentes",
            "Capítulo 4 - Os Independentes",
            [chapter4],
        ),
        "01_BACKGROUND_NARRADOR/faccoes/Segunda_Inquisicao.md": render_compilation(
            "Segunda Inquisição",
            "Capítulo 7 - seção 7.2",
            [sec_72],
        ),
        "01_BACKGROUND_NARRADOR/faccoes/index.md": build_faction_index(),
        "01_BACKGROUND_NARRADOR/index_personagens.md": render_compilation(
            "Índice de Personagens",
            "Capítulo 8 - seções 8.2, 8.3 e 8.4",
            [sec_82, sec_83, sec_84],
        ),
        "04_ANTAGONISTAS_V5/second_inquisition.md": render_compilation(
            "Second Inquisition",
            "Capítulo 7 - seção 7.2",
            [sec_72],
        ),
        "04_ANTAGONISTAS_V5/hunters.md": render_compilation(
            "Hunters",
            "Capítulo 7 - seção 7.3",
            [sec_73],
        ),
        "04_ANTAGONISTAS_V5/werewolves.md": render_compilation(
            "Werewolves",
            "Capítulo 7 - seção 7.4",
            [sec_74],
        ),
        "04_ANTAGONISTAS_V5/sabbat.md": render_compilation(
            "Sabbat",
            "Capítulo 7 - seção 7.5",
            [sec_75],
        ),
        "04_ANTAGONISTAS_V5/ghosts_and_occult.md": render_compilation(
            "Ghosts and Occult",
            "Capítulo 7 - seção 7.6",
            [sec_76],
        ),
        "04_ANTAGONISTAS_V5/cults.md": render_compilation(
            "Cults",
            "Capítulo 7 - seção 7.7",
            [sec_77],
        ),
        "04_ANTAGONISTAS_V5/creatures_index.md": render_compilation(
            "Creatures Index",
            "Capítulo 7 - Os Antagonistas",
            [chapter7],
        ),
    }

    entries_82 = collect_index_entries(sec_82, "8.2 - Índice completo por facção e clã")
    entries_83 = collect_index_entries(sec_83, "8.3 - Ghouls")
    essentials = collect_essentials(sec_84)

    general_docs["00_BACKGROUND_JOGADORES/rumores.md"] = build_rumors_doc(root, essentials)
    general_docs["01_BACKGROUND_NARRADOR/segredos_e_verdades.md"] = build_secrets_doc(essentials)
    general_docs["01_BACKGROUND_NARRADOR/teia_de_conexoes.md"] = build_teia_doc(
        entries_82,
        entries_83,
        essentials,
    )
    general_docs["01_BACKGROUND_NARRADOR/data/dominios.yml"] = build_dominios_yml(
        entries_82,
        entries_83,
        essentials,
    )
    general_docs["01_BACKGROUND_NARRADOR/data/rede_cainita.yml"] = build_rede_yml(
        entries_82,
        entries_83,
        essentials,
    )
    general_docs["01_BACKGROUND_NARRADOR/data/boon_e_dividas.yml"] = build_boons_yml(
        entries_82,
        essentials,
    )

    summary_files = list(MIRROR_ROOT.rglob("*_ficha_resumida.txt"))
    history_files = list(MIRROR_ROOT.rglob("*_historia.txt"))
    full_files = list(MIRROR_ROOT.rglob("*_ficha_completa.md"))
    prompt_files = list(PROMPT_MIRROR_DIR.glob("*.txt")) if PROMPT_MIRROR_DIR.exists() else []

    summary_lookup = build_lookup(summary_files, "_ficha_resumida")
    history_lookup = build_lookup(history_files, "_historia")
    full_lookup = build_lookup(full_files, "_ficha_completa")
    prompt_lookup = build_lookup(prompt_files, None)

    index_by_name = {entry.name: entry for entry in entries_82 + entries_83}

    filled_paths: set[Path] = set()
    missing_summary: list[str] = []
    missing_history: list[str] = []
    missing_full: list[str] = []
    missing_prompts: list[str] = []
    legacy_name_hits: list[str] = []

    for name, essential in sorted(essentials.items()):
        index_entry = index_by_name.get(name)

        summary_targets = resolve_paths(name, summary_lookup)
        if summary_targets:
            for path in summary_targets:
                write_text(path, format_essential_summary(essential, index_entry))
                filled_paths.add(path)
        else:
            missing_summary.append(name)

        history_targets = resolve_paths(name, history_lookup)
        if history_targets:
            for path in history_targets:
                write_text(path, format_essential_history(essential, index_entry))
                filled_paths.add(path)
        else:
            missing_history.append(name)

        full_targets = resolve_paths(name, full_lookup)
        if full_targets:
            for path in full_targets:
                write_text(path, format_essential_full(essential))
                filled_paths.add(path)
        else:
            missing_full.append(name)

        prompt_targets = resolve_paths(name, prompt_lookup)
        prompt_text = format_prompt_from_essential(essential, index_entry)
        if prompt_targets and prompt_text:
            for path in prompt_targets:
                write_text(path, prompt_text)
                filled_paths.add(path)
        elif not prompt_targets:
            missing_prompts.append(name)

        alias = LEGACY_FILE_ALIASES.get(norm_key(name))
        if alias:
            legacy_name_hits.append(f"{name} -> {alias}")

    for entry in sorted(entries_82 + entries_83, key=lambda item: item.name):
        if entry.name in essentials:
            continue

        summary_targets = resolve_paths(entry.name, summary_lookup)
        if summary_targets:
            for path in summary_targets:
                write_text(path, format_index_summary(entry))
                filled_paths.add(path)
        else:
            missing_summary.append(entry.name)

        history_targets = resolve_paths(entry.name, history_lookup)
        if history_targets:
            for path in history_targets:
                write_text(path, format_index_history(entry))
                filled_paths.add(path)
        else:
            missing_history.append(entry.name)

        alias = LEGACY_FILE_ALIASES.get(norm_key(entry.name))
        if alias:
            legacy_name_hits.append(f"{entry.name} -> {alias}")

    clan_entries: dict[str, list[IndexEntry]] = defaultdict(list)
    for entry in entries_82:
        if entry.classification:
            clan_entries[entry.classification].append(entry)

    for path in sorted((MIRROR_ROOT / "02_NPCS").rglob("estrutura_do_cla.md")):
        clan_label = slug_folder_label(path.parent.name)
        matches = [
            item
            for key, values in clan_entries.items()
            if norm_key(key, drop_articles=False) == norm_key(clan_label, drop_articles=False)
            or norm_key(key, drop_articles=False)
            == norm_key(clan_label.replace(" ", "-"), drop_articles=False)
            for item in values
        ]
        if matches:
            write_text(path, build_clan_structure(clan_label, matches))
            filled_paths.add(path)

    mortal_entries = [entry for entry in entries_82 if entry.section == "Segunda Inquisição"]
    general_docs["01_BACKGROUND_NARRADOR/faccoes/Mortal.md"] = build_mortal_faction_doc(mortal_entries)

    for relative_path, text in general_docs.items():
        path = MIRROR_ROOT / relative_path
        write_text(path, text)
        filled_paths.add(path)

    all_files = [path for path in MIRROR_ROOT.rglob("*") if path.is_file()]
    unmatched_files = sorted(
        str(path.relative_to(MIRROR_ROOT))
        for path in all_files
        if path.stat().st_size == 0
    )

    parse_failures = [entry for entry in entries_82 + entries_83 if entry.parse_failed]
    total_empty = len(unmatched_files)

    report_text = build_population_report(
        total_empty=total_empty,
        entries=entries_82,
        ghouls=entries_83,
        essentials=essentials,
        parse_failures=parse_failures,
        missing_summary=sorted(dict.fromkeys(missing_summary)),
        missing_history=sorted(dict.fromkeys(missing_history)),
        missing_full=sorted(dict.fromkeys(missing_full)),
        missing_prompts=sorted(dict.fromkeys(missing_prompts)),
        unmatched_files=unmatched_files,
        legacy_name_hits=sorted(dict.fromkeys(legacy_name_hits)),
    )
    write_text(MIRROR_ROOT / "01_BACKGROUND_NARRADOR/relatorio_populacao_e_lacunas.md", report_text)

    gap_text = build_gap_report(
        missing_summary=sorted(dict.fromkeys(missing_summary)),
        missing_history=sorted(dict.fromkeys(missing_history)),
        missing_full=sorted(dict.fromkeys(missing_full)),
        missing_prompts=sorted(dict.fromkeys(missing_prompts)),
        unmatched_files=unmatched_files,
        legacy_name_hits=sorted(dict.fromkeys(legacy_name_hits)),
    )
    write_text(MIRROR_ROOT / "LACUNAS_PARA_DECISAO.md", gap_text)

    mapping_lines = [
        "# Mapeamento do DOCX",
        "",
        "Reconstrução feita diretamente do arquivo `Sao_Paulo_by_Night.docx`.",
        "",
        "## Arquivos gerais preenchidos nesta passada",
        "",
    ]
    for relative_path in sorted(general_docs):
        mapping_lines.append(f"- `{relative_path}`")
    mapping_lines.extend(
        [
            "",
            "## Cobertura de personagens",
            "",
            f"- Entradas de índice 8.2: {len(entries_82)}",
            f"- Entradas de índice 8.3: {len(entries_83)}",
            f"- Fichas essenciais 8.4: {len(essentials)}",
            "",
            "## Observações",
            "",
            "- Fichas e histórias foram preenchidas só quando havia apoio direto no livro.",
            "- Arquivos que continuam vazios aparecem em `LACUNAS_PARA_DECISAO.md`.",
            "- Arquivos HTML/JS/CSS do site e do mapa ficaram fora desta passada de canonização de conteúdo.",
        ]
    )
    write_text(MIRROR_ROOT / "MAPEAMENTO_DOCX.md", tidy_lines(mapping_lines))


if __name__ == "__main__":
    main()
